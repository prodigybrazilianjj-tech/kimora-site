#!/usr/bin/env python3
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CREAM="F4EFE3"; CHAR="16140F"; CORAL="D8532E"; INK="1d1a14"; LINE="D8CFBA"; SHADE="FBF7EE"

doc=Document()
# default font
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(10.5)
st.font.color.rgb=RGBColor.from_string(INK)
for hid,sz in [("Heading 1",15),("Heading 2",12),("Heading 3",11)]:
    h=doc.styles[hid]; h.font.name="Arial"; h.font.size=Pt(sz); h.font.bold=True
    h.font.color.rgb=RGBColor.from_string(CHAR)

sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8)
sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def setcellfont(cell,bold=False,sz=10,color=INK):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name="Arial"; r.font.size=Pt(sz); r.font.bold=bold
            r.font.color.rgb=RGBColor.from_string(color)

def kvtable(rows,w0=2.3,w1=4.8):
    t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for k,v in rows:
        cells=t.add_row().cells
        cells[0].text=k; cells[1].text=v
        cells[0].width=Inches(w0); cells[1].width=Inches(w1)
        shade(cells[0],"EFE9DC"); setcellfont(cells[0],bold=True,sz=9.5,color=CHAR)
        setcellfont(cells[1],sz=9.5)
    return t

def para(text,bold=False,size=10.5,color=INK,space=6,align=None):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.name="Arial"; r.font.bold=bold; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    p.paragraph_format.space_after=Pt(space)
    if align: p.alignment=align
    return p

def bullets(items,style="List Bullet"):
    for it in items:
        p=doc.add_paragraph(style=style); r=p.add_run(it)
        r.font.name="Arial"; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(INK)
        p.paragraph_format.space_after=Pt(2)

# ---------- Header / logo ----------
try:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("logo.png", width=Inches(2.4))
except Exception as e:
    pass
para("Gym Counter Display — Fabrication Spec & Request for Quote", bold=True, size=16, color=CHAR, space=2, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Kimora Co.  ·  Single-serve creatine stick “grab station”  ·  Initial run: 50 units",
    size=10, color="5F5747", space=2, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Prepared "+datetime.date.today().strftime("%B %d, %Y")+"  ·  alex@kimoraco.com  ·  kimoraco.com  ·  Sedona, AZ",
    size=9, color="8A8170", space=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- 1 What ----------
doc.add_heading("1.  What we’re producing",level=1)
para("A countertop point-of-purchase (POP) display — a “grab station” — for jiu-jitsu / BJJ gym front desks. "
     "It has three integrated zones:", space=4)
bullets([
 "Header placard — printed brand card with the wordmark + bear-octopus crest, the $49.99 bag price, and a scannable QR to kimoraco.com.",
 "Stick tray — open tiered holder presenting single-serve creatine sticks face-out so the full stick is visible; restocked from the top.",
 "Cash honor box — a locked box with a coin/bill slot for $2 single-stick sales (honor system).",
])
para("Purpose: drive first-time product trial at the point of training. We are seeding 50 partner gyms.", space=8)

# ---------- 2 Qty/timeline ----------
doc.add_heading("2.  Quantity, timeline & budget",level=1)
kvtable([
 ("Initial order","50 units. Please also quote 100 and 250 for volume comparison."),
 ("Sample first","1 prototype/sample required for sign-off before the full run."),
 ("Timeline","Vendor selection + prototype now; full production timed to product launch (~Dec 2026)."),
 ("Target unit cost","~$50/unit goal. We will pay more for the custom-acrylic build if the quality justifies it."),
 ("Ship to","Kimora Co., PO Box 20024, Sedona, AZ 86341 (street address available for freight)."),
])
para("",space=2)

# ---------- 3 Build options ----------
doc.add_heading("3.  Two build options — please quote both",level=1)
doc.add_heading("Option A — Assembled kit  (target ~$50/unit)",level=3)
bullets([
 "Off-the-shelf locking acrylic donation/ballot box with coin slot + cam lock and 2 keys.",
 "Acrylic tiered bin / pack dispenser for the sticks (face-out).",
 "Rigid printed header (3–5 mm PVC or styrene) seated above/behind the tray.",
 "Full-color vinyl decal wrap on the cash box (bear-octopus crest + “Grab a stick — $2”).",
])
doc.add_heading("Option C — Custom-fabricated acrylic unit  (premium — quote actual)",level=3)
bullets([
 "Single integrated laser-cut / bonded acrylic unit (3–5 mm), matching the approved concept render.",
 "Direct-print or applied graphics on the header; integrated lockable coin slot in the base.",
 "Cleanest look and most durable; expected higher unit cost at 50 pcs — we want the real number.",
])

# ---------- 4 Specs ----------
doc.add_heading("4.  Specifications",level=1)
kvtable([
 ("Overall footprint","~10 in H × 6 in W × 4 in D (confirm / propose)"),
 ("Header placard","~6 in W × 4 in H, full-color (CMYK + white)"),
 ("Stick capacity","24–36 single-serve sticks, face-out, ~2 tiers; top restock"),
 ("Single stick (approx.)","~120 × 30 × 10 mm — exact dieline provided on award"),
 ("Cash box","Integrated, lockable (cam lock + 2 keys), coin + bill slot"),
 ("Material — Option A","Off-the-shelf acrylic box + PVC/styrene header + vinyl"),
 ("Material — Option C","Cast/extruded acrylic, 3–5 mm, bonded"),
 ("Print / finish","CMYK + white; UV-print or vinyl; scannable QR embedded as vector"),
 ("Brand colors","Cream #F4EFE3 · Charcoal #16140F · Coral #D8532E"),
])
para("",space=2)

# ---------- 5 Art ----------
doc.add_heading("5.  Art & files (provided on award)",level=1)
bullets([
 "Print-ready vector files (.ai / .eps / .pdf) for wordmark, stick and pouch — from our designer.",
 "Bear-octopus crest (the logo illustration). It MUST be reproduced exactly — the full octopus head (left) and snarling bear (right). No substitutions or AI redraws.",
 "Scannable QR to kimoraco.com, supplied as vector for the header.",
 "Approved concept render attached separately for visual reference (not production art).",
])

# ---------- 6 Questions ----------
doc.add_heading("6.  What we need back from you",level=1)
bullets([
 "Unit price at 50 / 100 / 250 units.",
 "One-time fees (tooling, setup, plates, die).",
 "Sample / prototype cost and lead time.",
 "Production lead time after art approval + deposit.",
 "Minimum order quantity (MOQ).",
 "Materials & thickness used; lock type and security.",
 "Freight cost to Sedona, AZ 86341, plus packed carton dimensions and weight.",
 "Payment terms.",
 "Print method, and confirmation the QR + fine crest detail hold at print size.",
])

# ---------- 7 Vendor shortlist ----------
doc.add_heading("7.  Vendor shortlist",level=1)
t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.LEFT
hdr=t.rows[0].cells
for i,htext in enumerate(["Vendor / path","Best for","Where"]):
    hdr[i].text=htext; shade(hdr[i],CHAR); setcellfont(hdr[i],bold=True,sz=9.5,color=CREAM)
rows=[
 ("ShopPOPdisplays (US)","Custom acrylic, faster turnaround","shoppopdisplays.com"),
 ("Alibaba acrylic POP suppliers","Custom acrylic, lowest unit cost at MOQ","alibaba.com — search “acrylic counter display”"),
 ("Local acrylic / laser shop","One-off prototype, fast iteration","Search “acrylic fabrication Phoenix/Sedona AZ”"),
 ("Displays2Go","Off-the-shelf locking boxes + bins","displays2go.com"),
 ("Source One Displays","Locking donation/ballot boxes","sourceone.org"),
 ("Uline / Amazon","Acrylic bins, locking boxes, fast","uline.com / amazon.com"),
 ("Local print shop","Rigid header + vinyl decal (Option A)","Search “sign shop / large-format print AZ”"),
]
for r0 in rows:
    cells=t.add_row().cells
    for i,val in enumerate(r0):
        cells[i].text=val; setcellfont(cells[i],sz=9)
    cells[0].width=Inches(2.4); cells[1].width=Inches(3.0); cells[2].width=Inches(2.6)
    shade(cells[0],SHADE)
para("",space=2)

# ---------- 8 Outreach template ----------
doc.add_heading("8.  Outreach email template",level=1)
para("Subject: Custom countertop retail display — quote request (50 units, repeatable)", bold=True, size=10, space=4)
body=("Hi [Vendor],\n\n"
"I’m sourcing a custom countertop point-of-purchase display for a creatine + electrolytes brand (Kimora Co.). "
"It’s a small “grab station” for gym front desks: a printed header placard, a tiered tray that holds single-serve "
"drink-mix sticks face-out, and a small locked cash box with a coin slot.\n\n"
"Initial run is 50 units, with more to follow if it performs — please quote 50 / 100 / 250. I’ve attached a spec "
"sheet and a concept render. Approx. footprint ~10″H × 6″W × 4″D; full-color header; integrated lock.\n\n"
"Could you send: unit price at each quantity, any one-time/tooling fees, sample cost + lead time, production lead time, "
"MOQ, materials, and freight to Sedona, AZ 86341? Print-ready vector art is ready on award.\n\n"
"Thanks,\nAlexzander Estrada\nFounder, Kimora Co.\nalex@kimoraco.com · kimoraco.com")
for line in body.split("\n"):
    para(line if line else " ", size=10, space=2)

doc.add_paragraph()
para("— Kimora Co.  ·  PO Box 20024, Sedona, AZ 86341  ·  alex@kimoraco.com  ·  kimoraco.com  ·  Grow Stronger. Think Sharper.",
    size=8.5, color="8A8170", space=0, align=WD_ALIGN_PARAGRAPH.CENTER)

out="Kimora_Gym_Display_Spec_and_RFQ.docx"
doc.save(out)
print("saved",out)
