#!/usr/bin/env python3
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHAR="16140F"; CORAL="D8532E"; INK="1d1a14"; SHADE="FBF7EE"; HDR="16140F"; CREAM="EFE9DC"

doc=Document()
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(10.5); st.font.color.rgb=RGBColor.from_string(INK)
for hid,sz in [("Heading 1",15),("Heading 2",12),("Heading 3",11)]:
    h=doc.styles[hid]; h.font.name="Arial"; h.font.size=Pt(sz); h.font.bold=True; h.font.color.rgb=RGBColor.from_string(CHAR)
sec=doc.sections[0]
sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:fill'),hexc); tcPr.append(s)
def cf(cell,bold=False,sz=9.5,color=INK):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name="Arial"; r.font.size=Pt(sz); r.font.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
def kv(rows,w0=2.4,w1=4.7):
    t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for k,v in rows:
        c=t.add_row().cells; c[0].text=k; c[1].text=v
        c[0].width=Inches(w0); c[1].width=Inches(w1)
        shade(c[0],CREAM); cf(c[0],bold=True,sz=9.5,color=CHAR); cf(c[1],sz=9.5)
    return t
def para(text,bold=False,size=10.5,color=INK,space=6,align=None):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.name="Arial"; r.font.bold=bold; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    p.paragraph_format.space_after=Pt(space)
    if align: p.alignment=align
    return p
def bullets(items):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); r=p.add_run(it)
        r.font.name="Arial"; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(INK)
        p.paragraph_format.space_after=Pt(2)

# Title
para("Gym Counter Display — Fabrication Spec & Quote Request", bold=True, size=16, color=CHAR, space=2, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Kimora Co.  ·  Custom acrylic countertop POP display  ·  Initial run: 50 units",
     size=10, color="5F5747", space=8, align=WD_ALIGN_PARAGRAPH.CENTER)
# Render
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("Kimora_Display_Render.png", width=Inches(6.3))
para("Approved design reference. Build to match this render.", size=8.5, color="8A8170", space=8, align=WD_ALIGN_PARAGRAPH.CENTER)

# Overview
doc.add_heading("1.  Overview",level=1)
para("A countertop point-of-purchase display for jiu-jitsu / BJJ gym front desks, built to match the render above. "
     "It is TWO pieces that sit side by side:", space=4)
bullets([
 "Module A — Header + stick tray: an open-front tray that presents single-serve creatine sticks standing upright and face-out, with a tall printed header panel rising behind it.",
 "Module B — Locking cash box: a separate closed box with a coin/bill slot and a keyed lock, for honor-system $2 single-stick sales.",
])
para("Purpose: first-time product trial at the point of training. Seeding 50 partner gyms.", space=6)

# Module A
doc.add_heading("2.  Module A — Header + stick tray",level=1)
bullets([
 "Open-front tray, matte black, with internal vertical channels that hold sticks upright, face-out. Render shows 6 facings; each channel holds several sticks deep (front-restock or top-restock).",
 "Header panel rises behind/above the tray at a slight backward angle. Cream background on the left ~2/3; charcoal panel on the right ~1/3 (behind the price + QR).",
 "Header graphics: KIMORA wordmark, the bear-octopus crest, “CREATINE + ELECTROLYTES,” the “Grow Stronger. Think Sharper.” tagline, a red diagonal “MORE FLAVORS COMING SOON” corner banner (top-left), a red “BAG PRICE $49.99” tag (top-right), a QR code, and “KIMORACO.COM.”",
 "Tray front rail printed: “STRAWBERRY GUAVA” (large, coral) with “STRENGTH* | HYDRATION* | FOCUS* | NO JUNK” beneath.",
])

# Module B
doc.add_heading("3.  Module B — Locking cash box",level=1)
bullets([
 "Separate closed box, matte black, roughly cube-shaped, slightly shorter than the header.",
 "Coin/bill slot on the top face; keyed cam lock (with 2 keys) on the upper-front face.",
 "Front graphics: “GRAB A STICK” / “$2” (large, coral) / “CASH HONOR SYSTEM,” with the bear-octopus emblem in coral at the bottom.",
 "Must be secure enough to deter casual tampering (locked rear or bottom access for cash retrieval).",
])

# Specs
doc.add_heading("4.  Specifications (approximate — please confirm / optimize)",level=1)
kv([
 ("Module A footprint","~10 in W × ~6 in H (tray) with header rising to ~9 in overall · ~4–5 in D"),
 ("Header panel","~10 in W × ~3 in H (≈3.3:1 wide banner, measured from render) · cream left / charcoal price panel right · full-color"),
 ("Stick facings / capacity","6 facings, face-out; ~24–36 sticks total (multiple deep per channel)"),
 ("Single stick size","~120 × 30 × 10 mm (≈4.7 × 1.2 × 0.4 in) — flat stick packet, not a bottle"),
 ("Module B (cash box)","~5 in W × ~6 in H × ~5 in D · coin/bill slot + keyed cam lock + 2 keys"),
 ("Material","Combination black + clear/printed acrylic, ~1/8 in (3 mm) — economical option confirmed by fabricator"),
 ("Print / finish","UV-printed graphics (header, tray rail, cash box); scannable QR embedded as vector"),
 ("Brand colors","Cream #F4EFE3 · Charcoal #16140F · Coral #D8532E"),
])
para("",space=2)

# Art
doc.add_heading("5.  Graphics & files (provided on award)",level=1)
bullets([
 "Print-ready vector files (.ai / .eps / .pdf) for the wordmark, crest, stick and pouch — from our designer.",
 "The bear-octopus crest MUST be reproduced exactly — full octopus head (left) and snarling bear (right). No substitutions.",
 "Scannable QR to kimoraco.com, supplied as vector for the header.",
 "This render is the design reference for layout, proportion, and finish (not production art).",
])

# Qty / timeline
doc.add_heading("6.  Quantity, timeline & budget",level=1)
kv([
 ("Initial order","50 units. Please also quote 100 and 250."),
 ("Sample first","1 prototype/sample for sign-off before the full run."),
 ("Timeline","Vendor + prototype now; full production timed to product launch (~Dec 2026)."),
 ("Target unit cost","~$50/unit goal; flexible upward if the build quality justifies it."),
 ("Ship to","Kimora Co., Sedona, AZ 86341 (street address for freight on request)."),
])
para("",space=2)

# RFQ
doc.add_heading("7.  What we need quoted",level=1)
bullets([
 "Unit price at 50 / 100 / 250 (Module A + Module B together).",
 "One-time fees (tooling, setup, print plates).",
 "Prototype / sample cost and lead time.",
 "Production lead time after art approval + deposit.",
 "Material & thickness used; lock type and cash-box security detail.",
 "Freight to Sedona, AZ 86341, with packed carton dimensions and weight.",
 "Payment terms; confirmation the QR and fine crest detail hold at print size.",
])

doc.add_paragraph()
para("— Kimora Co.  ·  alex@kimoraco.com  ·  kimoraco.com  ·  Sedona, Arizona  ·  Grow Stronger. Think Sharper.",
     size=8.5, color="8A8170", space=0, align=WD_ALIGN_PARAGRAPH.CENTER)

out="Kimora_Gym_Display_Spec_v2.docx"
doc.save(out)
print("saved",out)
